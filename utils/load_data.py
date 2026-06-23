# ───────────────────────────────────────────────────────────────────────────
#  load_data.py
# ───────────────────────────────────────────────────────────────────────────
import re
import logging
from pathlib import Path
from typing import Any, Optional
from langchain_core.documents import Document
import logfire
import numpy as np
from PIL import Image
from rich.console import Console
from tqdm import tqdm

# ── OCR stack ──────────────────────────────────────────────────────────────
try:
    import fitz                          # PyMuPDF
    import easyocr
    # GPU=False → CPU mode; set gpu=True if CUDA is available
    ocr_reader = easyocr.Reader(["en","fr"], gpu=False, verbose=False)
    OCR_AVAILABLE = True
except ImportError:
    fitz = None
    ocr_reader = None
    OCR_AVAILABLE = False
    logging.warning("PyMuPDF or EasyOCR not installed — OCR disabled.")

# ───────────────────────────────────────────────────────────────────────────
#  Bootstrap
# ───────────────────────────────────────────────────────────────────────────
console = Console()
logging.basicConfig(level=logging.WARNING)

# ───────────────────────────────────────────────────────────────────────────
#  Corpus loading & cleaning
# ───────────────────────────────────────────────────────────────────────────
 
def clean_reddit_text(text: str) -> str:
    """Strip Reddit PDF export noise — applied to PDF/OCR output only."""
    patterns = [
        r"Accéder au contenu principal.*?Se connecter",
        r"https?://\S+",
        r"\d{1,2}/\d{2,}",
        r"Sponsorisé\(e\).*?(?=\n[A-Z\[]|\Z)",
        r"(Répondre|Partager|Comm\. du top \d+%)",
        r"\[supprimé\]|\[removed\]",
        r"↑\s*-?\d+\s*↓|[⬆⬇]\s*\d+",
        r"Afficher plus de commentaires",
        r"\d+ réponses? supplémentaires?",
        r"Auteur[·\-]rice top \d+%",
        r"En savoir plus",
        r"Rechercher dans r/\w+",
        r"Trier par\s*:.*",
        r"Rechercher des commentaires",
    ]
    for pat in patterns:
        text = re.sub(pat, " ", text, flags=re.DOTALL | re.IGNORECASE)
    return re.sub(r"\s+", " ", text).strip()
 
 
def clean_generic_text(text: str) -> str:
    """Light cleaning for non-PDF sources (TXT, CSV, Excel, DOCX)."""
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\t", " ", text)
    text = re.sub(r" {3,}", "  ", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()
 
# tag each PDF with a topic before dropping it in ./corpus/.
# This metadata flows into ChunkMetadata and lets us slice RAGAS scores by topic in your comparison table
def infer_topic(filename: str) -> str:
    """Heuristic topic tag from filename — applied to all file types."""
    fname = filename.lower()
    if any(k in fname for k in ["playoff", "playoffs"]):        return "playoffs"
    if any(k in fname for k in ["stat", "stats"]):              return "player_stats"
    if any(k in fname for k in ["trade", "rumor"]):             return "trade_rumor"
    if any(k in fname for k in ["media", "rating", "snooze"]): return "media_bias"
    if any(k in fname for k in ["reddit"]):                     return "reddit_discussion"
    if any(k in fname for k in ["game", "recap", "score"]):    return "game_recap"
    return "general"
 
 
# ── PDF ──────────────────────────────────────────────────────────────────────
 
def extract_text_from_pdf_with_ocr(pdf_path: Path) -> Optional[list[tuple[int, str]]]:
    """
    Extract text from a PDF.
    Strategy: try PyPDF2 text layer first; if < 100 chars, fall back to EasyOCR.
    Returns list of (page_num, page_text) tuples, or None on total failure.
    """
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(pdf_path))
        pages_text = [(i, page.extract_text() or "") for i, page in enumerate(reader.pages)]
        total_chars = sum(len(t) for _, t in pages_text)
        if total_chars >= 100:
            logfire.info("PDF text via PyPDF2", file=pdf_path.name, chars=total_chars)
            return pages_text
        logfire.info("PyPDF2 found little text — falling back to OCR",
                     file=pdf_path.name, chars=total_chars)
    except Exception as e:
        logfire.warn("PyPDF2 failed", file=pdf_path.name, error=str(e))
 
    if not OCR_AVAILABLE:
        logfire.error("OCR not available — install pymupdf and easyocr")
        return None
 
    results: list[tuple[int, str]] = []
    try:
        doc = fitz.open(str(pdf_path))
        for page_num in tqdm(range(len(doc)), desc=f"OCR {pdf_path.name}", leave=False):
            page = doc.load_page(page_num)
            pix  = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img  = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            try:
                ocr_results = ocr_reader.readtext(np.array(img))
                page_text   = "\n".join(res[1] for res in ocr_results if res[2] >= 0.3)
                results.append((page_num, page_text))
            except Exception as ocr_err:
                logfire.warn("OCR failed on page", file=pdf_path.name,
                             page=page_num + 1, error=str(ocr_err))
        doc.close()
        return results if results else None
    except Exception as e:
        logfire.error("PDF open failed", file=str(pdf_path), error=str(e))
        return None
 
 
def _pdf_to_documents(pdf_path: Path) -> list[Document]:
    """Convert a PDF into one Document per page after cleaning."""
    page_results = extract_text_from_pdf_with_ocr(pdf_path)
    if not page_results:
        logfire.warn("No text extracted from PDF", file=pdf_path.name)
        return []
    topic = infer_topic(pdf_path.stem)
    docs  = []
    for page_num, raw_text in page_results:
        cleaned = clean_reddit_text(raw_text)
        if len(cleaned) < 80:
            continue
        docs.append(Document(
            page_content=cleaned,
            metadata={"source": pdf_path.stem, "page": page_num, "topic": topic,
                      "file_type": "pdf", "file_path": str(pdf_path)},
        ))
    logfire.info("PDF loaded", file=pdf_path.name,
                 total_pages=len(page_results), kept_pages=len(docs))
    return docs
 
 
# ── TXT ──────────────────────────────────────────────────────────────────────
 
def _txt_to_documents(txt_path: Path) -> list[Document]:
    """Load a plain-text file as a single Document. Tries UTF-8 then latin-1."""
    text = None
    for enc in ("utf-8", "latin-1"):
        try:
            text = txt_path.read_text(encoding=enc, errors="ignore")
            break
        except Exception:
            continue
    if not text or len(text.strip()) < 20:
        logfire.warn("TXT empty or too short", file=txt_path.name)
        return []
    cleaned = clean_generic_text(text)
    logfire.info("TXT loaded", file=txt_path.name, chars=len(cleaned))
    return [Document(
        page_content=cleaned,
        metadata={"source": txt_path.stem, "page": 0, "topic": infer_topic(txt_path.stem),
                  "file_type": "txt", "file_path": str(txt_path)},
    )]
 
 
# ── CSV ──────────────────────────────────────────────────────────────────────
 
def _csv_to_documents(csv_path: Path) -> list[Document]:
    """
    Load a CSV as a single Document.
    Tries UTF-8, latin-1, semicolon separator — matching the prototype's fallback chain.
    Each row becomes "col1: val1 | col2: val2 | ..." for retrieval-friendly format.
    """
    import pandas as pd
    df = None
    for kwargs in [{"encoding": "utf-8"}, {"encoding": "latin-1"},
                   {"sep": ";", "encoding": "utf-8"}, {"sep": ";", "encoding": "latin-1"}]:
        try:
            df = pd.read_csv(csv_path, **kwargs)
            break
        except Exception:
            continue
    if df is None or df.empty:
        logfire.warn("CSV unreadable or empty", file=csv_path.name)
        return []
    rows = [
        " | ".join(
            f"{col}: {val}" for col, val in row.items()
            if pd.notna(val) and str(val).strip()
        )
        for _, row in df.iterrows()
    ]
    rows = [r for r in rows if r]
    if not rows:
        return []
    text    = f"File: {csv_path.name}\nColumns: {', '.join(df.columns)}\n\n" + "\n".join(rows)
    cleaned = clean_generic_text(text)
    logfire.info("CSV loaded", file=csv_path.name, rows=len(rows), chars=len(cleaned))
    return [Document(
        page_content=cleaned,
        metadata={"source": csv_path.stem, "page": 0, "topic": infer_topic(csv_path.stem),
                  "file_type": "csv", "file_path": str(csv_path),
                  "row_count": len(rows), "columns": list(df.columns)},
    )]
 
 
# ── EXCEL ─────────────────────────────────────────────────────────────────────
 
def _excel_to_documents(excel_path: Path) -> list[Document]:
    """
    Load Excel as one Document per sheet.
    Multi-sheet workbooks keep each sheet separate for retrieval granularity.
    """
    import pandas as pd
    try:
        xl = pd.ExcelFile(excel_path)
    except Exception as e:
        logfire.error("Excel open failed", file=excel_path.name, error=str(e))
        return []
    docs  = []
    topic = infer_topic(excel_path.stem)
    for idx, sheet_name in enumerate(xl.sheet_names):
        try:
            df = xl.parse(sheet_name)
            if df.empty:
                continue
            rows = [
                " | ".join(
                    f"{col}: {val}" for col, val in row.items()
                    if pd.notna(val) and str(val).strip()
                )
                for _, row in df.iterrows()
            ]
            rows = [r for r in rows if r]
            if not rows:
                continue
            text = (
                f"File: {excel_path.name} | Sheet: {sheet_name}\n"
                f"Columns: {', '.join(str(c) for c in df.columns)}\n\n"
                + "\n".join(rows)
            )
            cleaned = clean_generic_text(text)
            docs.append(Document(
                page_content=cleaned,
                metadata={"source": excel_path.stem, "page": idx,
                          "sheet_name": sheet_name, "topic": topic,
                          "file_type": "excel", "file_path": str(excel_path),
                          "row_count": len(rows), "columns": [str(c) for c in df.columns]},
            ))
            logfire.info("Excel sheet loaded", file=excel_path.name,
                         sheet=sheet_name, rows=len(rows))
        except Exception as e:
            logfire.warn("Excel sheet failed", file=excel_path.name,
                         sheet=sheet_name, error=str(e))
    logfire.info("Excel loaded", file=excel_path.name,
                 sheets_total=len(xl.sheet_names), sheets_kept=len(docs))
    return docs
 
 
# ── DOCX ──────────────────────────────────────────────────────────────────────
 
def _docx_to_documents(docx_path: Path) -> list[Document]:
    """Load a Word document — extracts paragraphs only."""
    try:
        import docx as python_docx
        doc  = python_docx.Document(str(docx_path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logfire.error("DOCX extraction failed", file=docx_path.name, error=str(e))
        return []
    if len(text.strip()) < 20:
        logfire.warn("DOCX empty or too short", file=docx_path.name)
        return []
    cleaned = clean_generic_text(text)
    logfire.info("DOCX loaded", file=docx_path.name, chars=len(cleaned))
    return [Document(
        page_content=cleaned,
        metadata={"source": docx_path.stem, "page": 0, "topic": infer_topic(docx_path.stem),
                  "file_type": "docx", "file_path": str(docx_path)},
    )]
 
 
# ── Dispatcher ────────────────────────────────────────────────────────────────
 
_LOADERS: dict[str, Any] = {
    ".pdf":  _pdf_to_documents,
    ".txt":  _txt_to_documents,
    ".csv":  _csv_to_documents,
    ".xlsx": _excel_to_documents,
    ".xls":  _excel_to_documents,
    ".docx": _docx_to_documents,
}
 
 
@logfire.instrument("load_corpus")
def load_corpus(corpus_dir: Path) -> list[Document]:
    """
    Load all supported files from corpus_dir.
    Supported: .pdf · .txt · .csv · .xlsx / .xls · .docx
    All loaders return list[Document] with consistent metadata keys:
      source, page, topic, file_type, file_path
    """
    docs:       list[Document] = []
    file_count: dict[str, int] = {}
 
    all_files = [
        f for f in corpus_dir.rglob("*")
        if f.is_file() and f.suffix.lower() in _LOADERS
    ]
    if not all_files:
        logfire.warn("No supported files found", path=str(corpus_dir),
                     supported=list(_LOADERS.keys()))
        raise FileNotFoundError(
            f"No supported files in {corpus_dir}. Supported: {list(_LOADERS.keys())}"
        )
 
    for file_path in all_files:
        ext    = file_path.suffix.lower()
        loader = _LOADERS[ext]
        with logfire.span("load_file", file=file_path.name, type=ext):
            try:
                file_docs = loader(file_path)
                docs.extend(file_docs)
                file_count[ext] = file_count.get(ext, 0) + 1
            except Exception as exc:
                logfire.error("File load failed — skipped",
                              file=file_path.name, error=str(exc))
 
    logfire.info("Corpus loaded", total_docs=len(docs), files_by_type=file_count)
    console.print(
        f"  [green]Corpus:[/green] {len(docs)} documents from "
        + ", ".join(f"{v} {k}" for k, v in file_count.items())
    )
    return docs